import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from PyPDF2 import PdfWriter, PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

DEMO_DATA_DIR = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(DEMO_DATA_DIR, "docs")
os.makedirs(DOCS_DIR, exist_ok=True)


def create_work_weekly_report():
    doc = Document()
    
    title = doc.add_heading('工作周报', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('一、本周工作总结', level=2)
    
    doc.add_paragraph('1. 项目进度')
    doc.add_paragraph('   - 完成用户管理模块的前端页面开发')
    doc.add_paragraph('   - 编写API接口文档，完成接口联调')
    doc.add_paragraph('   - 修复了3个线上bug')
    
    doc.add_paragraph('2. 团队协作')
    doc.add_paragraph('   - 参加了2次需求评审会议')
    doc.add_paragraph('   - 帮助新人熟悉项目架构')
    doc.add_paragraph('   - 完成代码review任务')
    
    doc.add_heading('二、下周工作计划', level=2)
    doc.add_paragraph('   - 完成订单管理模块开发')
    doc.add_paragraph('   - 编写单元测试用例')
    doc.add_paragraph('   - 准备技术分享')
    
    doc.add_heading('三、问题与建议', level=2)
    doc.add_paragraph('   - 测试环境不稳定，建议增加自动化测试')
    doc.add_paragraph('   - 文档更新不及时，影响开发效率')
    
    doc.add_paragraph('')
    doc.add_paragraph('汇报人：张三', style='Normal')
    doc.add_paragraph('日期：2024年3月15日', style='Normal')
    
    doc.save(os.path.join(DOCS_DIR, '工作周报.docx'))
    print('Created: 工作周报.docx')


def create_study_notes():
    doc = Document()
    
    title = doc.add_heading('Python数据分析学习笔记', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('第一章：NumPy基础', level=2)
    
    doc.add_paragraph('1.1 数组创建')
    doc.add_paragraph('   - numpy.array(): 创建数组')
    doc.add_paragraph('   - numpy.zeros(): 创建全零数组')
    doc.add_paragraph('   - numpy.ones(): 创建全一数组')
    doc.add_paragraph('   - numpy.arange(): 创建等差数列')
    
    doc.add_paragraph('1.2 数组运算')
    doc.add_paragraph('   - 元素级运算：+ - * /')
    doc.add_paragraph('   - 矩阵乘法：@ 或 dot()')
    doc.add_paragraph('   - 统计函数：mean(), sum(), max(), min()')
    
    doc.add_heading('第二章：Pandas数据处理', level=2)
    
    doc.add_paragraph('2.1 DataFrame基本操作')
    doc.add_paragraph('   - 读取数据：pd.read_csv(), pd.read_excel()')
    doc.add_paragraph('   - 查看数据：head(), tail(), info()')
    doc.add_paragraph('   - 筛选数据：loc[], iloc[]')
    
    doc.add_paragraph('2.2 数据清洗')
    doc.add_paragraph('   - 处理缺失值：dropna(), fillna()')
    doc.add_paragraph('   - 数据类型转换：astype()')
    doc.add_paragraph('   - 去重：drop_duplicates()')
    
    doc.add_paragraph('')
    doc.add_paragraph('学习日期：2024年3月10日', style='Normal')
    doc.add_paragraph('学习时长：3小时', style='Normal')
    
    doc.save(os.path.join(DOCS_DIR, '学习笔记.docx'))
    print('Created: 学习笔记.docx')


def create_diary():
    doc = Document()
    
    title = doc.add_heading('日记', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('2024年3月12日 星期三 晴', level=2)
    
    doc.add_paragraph('今天是个美好的日子。早上起来阳光明媚，心情特别好。')
    doc.add_paragraph('上午去公司处理了积压的工作，效率很高，中午前就完成了。')
    doc.add_paragraph('下午和同事一起去咖啡馆聊天，聊了很多关于生活和未来的话题。')
    doc.add_paragraph('晚上回家做了一顿丰盛的晚餐，犒劳一下自己。')
    doc.add_paragraph('最近感觉自己状态不错，希望能保持这种积极的心态。')
    
    doc.add_heading('2024年3月13日 星期四 多云', level=2)
    
    doc.add_paragraph('今天工作遇到了一些挑战，有个bug困扰了我很久。')
    doc.add_paragraph('不过最后还是解决了，那种成就感真的很棒。')
    doc.add_paragraph('下班后去健身房锻炼了一小时，出了一身汗感觉很舒服。')
    doc.add_paragraph('回家路上看到了美丽的晚霞，忍不住停下脚步欣赏了一会儿。')
    
    doc.save(os.path.join(DOCS_DIR, '日记.docx'))
    print('Created: 日记.docx')


def create_travel_notes():
    doc = Document()
    
    title = doc.add_heading('云南旅行日记', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('第一天：抵达昆明', level=2)
    
    doc.add_paragraph('上午10点到达昆明长水机场，天气比想象中暖和。')
    doc.add_paragraph('入住酒店后去了滇池，看到了很多红嘴鸥。')
    doc.add_paragraph('晚上吃了正宗的过桥米线，味道很不错。')
    
    doc.add_heading('第二天：大理古城', level=2)
    
    doc.add_paragraph('早上坐高铁到大理，车程约2小时。')
    doc.add_paragraph('下午逛了大理古城，感受到了浓厚的历史氛围。')
    doc.add_paragraph('晚上在洋人街吃了小吃，很热闹。')
    
    doc.add_heading('第三天：洱海骑行', level=2)
    
    doc.add_paragraph('租了一辆电动车环洱海骑行，风景太美了！')
    doc.add_paragraph('一路上走走停停，拍了很多照片。')
    doc.add_paragraph('晚上住海景客栈，听着海浪声入睡。')
    
    doc.add_paragraph('')
    doc.add_paragraph('旅行心得：云南真的很美，下次还想来！', style='Normal')
    
    doc.save(os.path.join(DOCS_DIR, '旅行记录.docx'))
    print('Created: 旅行记录.docx')


def create_reading_notes():
    doc = Document()
    
    title = doc.add_heading('《人类简史》读书笔记', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('核心观点', level=2)
    
    doc.add_paragraph('1. 认知革命：大约7万年前，智人发展出了虚构故事的能力。')
    doc.add_paragraph('2. 农业革命：作者称之为"史上最大骗局"，人类从采集转向农耕。')
    doc.add_paragraph('3. 科学革命：承认自己的无知，通过观察和实验获取知识。')
    
    doc.add_heading('精彩段落', level=2)
    
    doc.add_paragraph('"农业革命所带来的非但不是轻松生活的新时代，反而让农民过着比采集者更辛苦、更不满足的生活。"')
    
    doc.add_paragraph('"人类凭借想象构建的秩序改变了现实世界。"')
    
    doc.add_heading('读后感', level=2)
    
    doc.add_paragraph('这本书让我重新思考了人类的历史和现状。很多我们习以为常的东西，')
    doc.add_paragraph('其实都是虚构的故事。比如国家、金钱、宗教，这些都是人类共同想象的产物。')
    doc.add_paragraph('读完这本书后，我看待世界的方式发生了改变。')
    
    doc.add_paragraph('')
    doc.add_paragraph('阅读日期：2024年2月', style='Normal')
    doc.add_paragraph('推荐指数：★★★★★', style='Normal')
    
    doc.save(os.path.join(DOCS_DIR, '读书笔记.docx'))
    print('Created: 读书笔记.docx')


def create_pdf_from_text(text_content, filename):
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    
    text = p.beginText(72, 750)
    text.setFont("Helvetica", 12)
    
    lines = text_content.split('\n')
    for line in lines:
        text.textLine(line)
    
    p.drawText(text)
    p.showPage()
    p.save()
    
    buffer.seek(0)
    with open(os.path.join(DOCS_DIR, filename), 'wb') as f:
        f.write(buffer.getvalue())
    print(f'Created: {filename}')


if __name__ == '__main__':
    create_work_weekly_report()
    create_study_notes()
    create_diary()
    create_travel_notes()
    create_reading_notes()
    
    create_pdf_from_text("""工作周报

一、本周工作总结
1. 项目进度
   - 完成用户管理模块的前端页面开发
   - 编写API接口文档，完成接口联调
   - 修复了3个线上bug

2. 团队协作
   - 参加了2次需求评审会议
   - 帮助新人熟悉项目架构
   - 完成代码review任务

二、下周工作计划
   - 完成订单管理模块开发
   - 编写单元测试用例
   - 准备技术分享

三、问题与建议
   - 测试环境不稳定，建议增加自动化测试
   - 文档更新不及时，影响开发效率

汇报人：张三
日期：2024年3月15日
""", '工作周报.pdf')
    
    create_pdf_from_text("""学习笔记

第一章：NumPy基础
1.1 数组创建
   - numpy.array(): 创建数组
   - numpy.zeros(): 创建全零数组
   - numpy.ones(): 创建全一数组
   - numpy.arange(): 创建等差数列

1.2 数组运算
   - 元素级运算：+ - * /
   - 矩阵乘法：@ 或 dot()
   - 统计函数：mean(), sum(), max(), min()

第二章：Pandas数据处理
2.1 DataFrame基本操作
   - 读取数据：pd.read_csv(), pd.read_excel()
   - 查看数据：head(), tail(), info()
   - 筛选数据：loc[], iloc[]

2.2 数据清洗
   - 处理缺失值：dropna(), fillna()
   - 数据类型转换：astype()
   - 去重：drop_duplicates()

学习日期：2024年3月10日
学习时长：3小时
""", '学习笔记.pdf')
    
    print('\nAll demo documents created successfully!')
